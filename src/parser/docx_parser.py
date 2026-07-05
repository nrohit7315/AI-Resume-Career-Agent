"""DOCX / TXT resume text extraction."""
from __future__ import annotations

from pathlib import Path

from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxParseError(Exception):
    """Raised when a DOCX/TXT file cannot be read."""


def extract_text_from_docx(file_path: str | Path) -> str:
    file_path = Path(file_path)
    if not file_path.exists():
        raise DocxParseError(f"File not found: {file_path}")

    try:
        import docx  # python-docx
    except ImportError as exc:
        raise DocxParseError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    try:
        document = docx.Document(str(file_path))
        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

        # Tables often hold skills/education in resumes built from templates.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise DocxParseError(f"Failed to parse DOCX '{file_path.name}': {exc}") from exc


def extract_text_from_txt(file_path: str | Path) -> str:
    file_path = Path(file_path)
    if not file_path.exists():
        raise DocxParseError(f"File not found: {file_path}")

    encodings = ["utf-8", "utf-8-sig", "latin-1"]
    last_err: Exception | None = None
    for enc in encodings:
        try:
            return file_path.read_text(encoding=enc)
        except Exception as exc:  # noqa: BLE001
            last_err = exc
            continue
    raise DocxParseError(f"Failed to read TXT '{file_path.name}': {last_err}")
